"""
Extended API routes for ResearchAI platform.

Includes:
- PDF paper upload and ingestion
- Visual diagram generation
- Academic structure formatting
- Enhanced export with Q1 journal formatting
"""

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, Response
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
import logging
import io

logger = logging.getLogger(__name__)

# Create router
extended_router = APIRouter(prefix="/api/v2", tags=["Extended Features"])


# ============================================================================
# Request/Response Models
# ============================================================================

class PDFUploadResponse(BaseModel):
    success: bool
    filename: str
    file_hash: str
    title: str
    authors: List[str]
    abstract: str
    sections_found: List[str]
    message: str


class DiagramRequest(BaseModel):
    topic: str
    methodology: Optional[str] = ""
    diagram_types: List[str] = ["workflow", "flowchart", "conceptual"]
    custom_requirements: List[str] = []


class DiagramResponse(BaseModel):
    diagrams: List[Dict[str, Any]]
    count: int


class AcademicFormatRequest(BaseModel):
    proposal_id: str
    author_name: str
    institution: Optional[str] = ""
    department: Optional[str] = ""
    supervisor_name: Optional[str] = ""
    dedication_to: Optional[str] = ""
    include_visuals: bool = True


# ============================================================================
# PDF Upload & Ingestion Routes
# ============================================================================

@extended_router.post("/papers/upload", response_model=PDFUploadResponse)
async def upload_pdf_paper(
    file: UploadFile = File(...),
    parse_immediately: bool = Form(default=True),
):
    """
    Upload a PDF research paper for literature review augmentation.
    
    The paper will be parsed to extract:
    - Title and authors
    - Abstract
    - Methodology sections
    - Key findings
    - References
    
    These will be used to enhance the literature review process.
    """
    try:
        # Validate file type
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(
                status_code=400,
                detail="Only PDF files are accepted"
            )
        
        # Import PDF service
        from src.services.pdf_ingestion import get_pdf_service
        
        pdf_service = get_pdf_service()
        
        # Read file content
        content = await file.read()
        
        if len(content) > 50 * 1024 * 1024:  # 50MB limit
            raise HTTPException(
                status_code=413,
                detail="File too large. Maximum size is 50MB."
            )
        
        # Save file
        file_path = await pdf_service.save_uploaded_file(file.filename, content)
        
        # Parse immediately if requested
        parsed_paper = None
        if parse_immediately:
            parsed_paper = await pdf_service.parse_pdf(file_path)
        
        if parsed_paper:
            return PDFUploadResponse(
                success=True,
                filename=file.filename,
                file_hash=parsed_paper.file_hash,
                title=parsed_paper.title or "Unknown Title",
                authors=parsed_paper.authors or [],
                abstract=parsed_paper.abstract[:500] if parsed_paper.abstract else "",
                sections_found=[
                    s for s in ["abstract", "introduction", "methodology", "results", "conclusion"]
                    if getattr(parsed_paper, s, None)
                ],
                message="Paper uploaded and parsed successfully",
            )
        else:
            return PDFUploadResponse(
                success=True,
                filename=file.filename,
                file_hash="",
                title="",
                authors=[],
                abstract="",
                sections_found=[],
                message="Paper uploaded. Parsing will occur during proposal generation.",
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@extended_router.get("/papers")
async def list_uploaded_papers():
    """List all uploaded PDF papers."""
    try:
        from src.services.pdf_ingestion import get_pdf_service
        
        pdf_service = get_pdf_service()
        papers = pdf_service.list_uploaded_papers()
        
        return {
            "papers": [
                {
                    "filename": p.name,
                    "path": str(p),
                    "size_kb": p.stat().st_size // 1024,
                }
                for p in papers
            ],
            "count": len(papers),
        }
    except Exception as e:
        logger.error(f"List papers error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@extended_router.post("/papers/parse-all")
async def parse_all_papers():
    """Parse all uploaded PDF papers."""
    try:
        from src.services.pdf_ingestion import get_pdf_service
        
        pdf_service = get_pdf_service()
        papers = await pdf_service.ingest_all_papers()
        
        return {
            "success": True,
            "papers_parsed": len(papers),
            "papers": [
                {
                    "title": p.title,
                    "authors": p.authors[:3],
                    "year": p.year,
                    "has_abstract": bool(p.abstract),
                    "has_methodology": bool(p.methodology),
                }
                for p in papers
            ],
        }
    except Exception as e:
        logger.error(f"Parse all papers error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Visual Diagram Generation Routes
# ============================================================================

@extended_router.post("/diagrams/generate", response_model=DiagramResponse)
async def generate_diagrams(request: DiagramRequest):
    """
    Generate visual diagrams for the research proposal.
    
    Available diagram types:
    - workflow: Research workflow diagram
    - flowchart: Methodology flowchart
    - conceptual: Conceptual framework
    - dataflow: Data flow diagram
    - architecture: System architecture
    - gantt: Project timeline
    """
    try:
        from src.agents.visual.visual_workflow_agent import VisualWorkflowAgent
        from src.models.agent_messages import AgentRequest
        from uuid import uuid4
        
        agent = VisualWorkflowAgent()
        
        agent_request = AgentRequest(
            task_id=str(uuid4()),
            agent_name="visual_workflow_agent",
            input_data={
                "topic": request.topic,
                "methodology": request.methodology,
                "diagram_types": request.diagram_types,
                "custom_requirements": request.custom_requirements,
            }
        )
        
        response = await agent.execute(agent_request)
        
        if response.status.value == "completed":
            return DiagramResponse(
                diagrams=response.output_data.get("diagrams", []),
                count=response.output_data.get("count", 0),
            )
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Diagram generation failed: {response.error}"
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Diagram generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@extended_router.get("/diagrams/types")
async def list_diagram_types():
    """List available diagram types."""
    return {
        "types": [
            {"id": "workflow", "name": "Research Workflow", "description": "Overview of research process steps"},
            {"id": "flowchart", "name": "Methodology Flowchart", "description": "Detailed methodology steps"},
            {"id": "conceptual", "name": "Conceptual Framework", "description": "Variables and relationships"},
            {"id": "dataflow", "name": "Data Flow Diagram", "description": "Data processing pipeline"},
            {"id": "architecture", "name": "System Architecture", "description": "Proposed system design"},
            {"id": "gantt", "name": "Project Timeline", "description": "Gantt chart for planning"},
        ]
    }


# ============================================================================
# Academic Structure Formatting Routes
# ============================================================================

@extended_router.post("/proposals/{proposal_id}/format-academic")
async def format_proposal_academic(proposal_id: str, request: AcademicFormatRequest):
    """
    Apply Q1/Scopus/Elsevier academic formatting to a proposal.
    
    This adds:
    - Title page
    - Dedication page
    - Acknowledgements
    - Table of contents
    - Proper chapter structure
    - Consolidated references
    """
    try:
        from src.services.academic_structure import create_academic_template
        from src.services.formatting_controller import get_formatting_controller
        
        # Get the proposal
        # Import from main module
        import sys
        main_module = sys.modules.get('src.api.main')
        if main_module:
            completed_proposals = getattr(main_module, 'completed_proposals', {})
            proposal = completed_proposals.get(proposal_id)
        else:
            proposal = None
        
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        # Create academic template
        template = create_academic_template(
            title=proposal.get("topic", "Research Proposal"),
            author_name=request.author_name,
            institution=request.institution,
            department=request.department,
            supervisor_name=request.supervisor_name,
            citation_style="harvard",
        )
        
        # Get full structure
        structure = template.get_full_structure()
        
        # Apply formatting controller
        formatter = get_formatting_controller()
        formatted_proposal = formatter.format_proposal(proposal)
        
        # Map content to structure
        structured = template.map_content_to_structure(formatted_proposal)
        
        return {
            "success": True,
            "proposal_id": proposal_id,
            "structure": structure,
            "formatted_proposal": formatted_proposal,
            "message": "Academic formatting applied successfully",
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Academic formatting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@extended_router.get("/proposals/{proposal_id}/export-academic/{format}")
async def export_academic_format(
    proposal_id: str,
    format: str,
    author_name: str = "Author Name",
    institution: str = "",
    supervisor_name: str = "",
):
    """
    Export proposal in full Q1 academic format.
    
    Formats: pdf, docx
    
    Includes all front matter (title page, dedication, acknowledgements, TOC)
    and proper chapter structure.
    """
    try:
        from src.services.academic_structure import create_academic_template, ProposalMetadata
        from src.services.formatting_controller import get_formatting_controller
        
        # Get the proposal
        import sys
        main_module = sys.modules.get('src.api.main')
        if main_module:
            completed_proposals = getattr(main_module, 'completed_proposals', {})
            proposal = completed_proposals.get(proposal_id)
        else:
            proposal = None
        
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        # Apply formatting
        formatter = get_formatting_controller()
        formatted = formatter.format_for_export(proposal, format)
        
        # Create academic template
        template = create_academic_template(
            title=proposal.get("topic", "Research Proposal"),
            author_name=author_name,
            institution=institution,
            supervisor_name=supervisor_name,
        )
        
        if format == "docx":
            return await _export_academic_docx(formatted, template, proposal_id)
        elif format == "pdf":
            return await _export_academic_pdf(formatted, template, proposal_id)
        else:
            raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Academic export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


async def _export_academic_docx(proposal: Dict, template: Any, proposal_id: str) -> Response:
    """Generate academic DOCX with full structure."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        document = Document()
        
        # Set margins
        for section in document.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(2.54)
        
        # ===== TITLE PAGE =====
        title_page = template.generate_title_page()
        
        # Add spacing before title
        for _ in range(5):
            document.add_paragraph()
        
        title = document.add_heading(title_page["content"]["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(3):
            document.add_paragraph()
        
        author = document.add_paragraph(title_page["content"]["author"])
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if title_page["content"]["institution"]:
            inst = document.add_paragraph(title_page["content"]["institution"])
            inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(5):
            document.add_paragraph()
        
        date = document.add_paragraph(title_page["content"]["date"])
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_page_break()
        
        # ===== DEDICATION PAGE =====
        dedication = template.generate_dedication_page()
        document.add_heading("DEDICATION", 1)
        document.add_paragraph(dedication["content"])
        document.add_page_break()
        
        # ===== ACKNOWLEDGEMENTS =====
        ack = template.generate_acknowledgement_page()
        document.add_heading("ACKNOWLEDGEMENTS", 1)
        for para in ack["content"].split("\n\n"):
            if para.strip():
                p = document.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_page_break()
        
        # ===== ABSTRACT =====
        document.add_heading("ABSTRACT", 1)
        
        # Get abstract from proposal sections
        sections = proposal.get("full_sections", [])
        abstract_content = ""
        for section in sections:
            if "introduction" in section.get("title", "").lower():
                # Use first 300 words of introduction as abstract basis
                words = section.get("content", "").split()[:300]
                abstract_content = " ".join(words) + "..."
                break
        
        if abstract_content:
            p = document.add_paragraph(abstract_content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        document.add_page_break()
        
        # ===== TABLE OF CONTENTS =====
        document.add_heading("TABLE OF CONTENTS", 1)
        toc_entries = template.generate_table_of_contents()
        
        for entry in toc_entries:
            level = entry.level
            indent = "    " * level
            p = document.add_paragraph(f"{indent}{entry.title}")
        
        document.add_page_break()
        
        # ===== MAIN CONTENT =====
        for i, section in enumerate(sections, 1):
            title = section.get("title", f"Section {i}")
            content = section.get("content", "")
            
            # Add chapter heading
            document.add_heading(title, 1)
            
            # Add content paragraphs
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    p = document.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Save to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{proposal_id[:8]}_academic_proposal.docx"'
            }
        )
        
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")


async def _export_academic_pdf(proposal: Dict, template: Any, proposal_id: str) -> Response:
    """Generate academic PDF with full structure."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
        
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2.54*cm,
            leftMargin=3.17*cm,
            topMargin=2.54*cm,
            bottomMargin=2.54*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Title'],
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30,
            fontName='Times-Bold'
        )
        
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=15,
            fontName='Times-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            firstLineIndent=1.27*cm,
        )
        
        center_style = ParagraphStyle(
            'CenterStyle',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            fontName='Times-Roman'
        )
        
        story = []
        
        # ===== TITLE PAGE =====
        title_page = template.generate_title_page()
        
        story.append(Spacer(1, 150))
        story.append(Paragraph(title_page["content"]["title"], title_style))
        story.append(Spacer(1, 100))
        story.append(Paragraph(title_page["content"]["author"], center_style))
        
        if title_page["content"]["institution"]:
            story.append(Spacer(1, 20))
            story.append(Paragraph(title_page["content"]["institution"], center_style))
        
        story.append(Spacer(1, 150))
        story.append(Paragraph(title_page["content"]["date"], center_style))
        story.append(PageBreak())
        
        # ===== DEDICATION =====
        dedication = template.generate_dedication_page()
        story.append(Paragraph("DEDICATION", heading_style))
        story.append(Spacer(1, 20))
        story.append(Paragraph(dedication["content"].replace("\n", "<br/>"), body_style))
        story.append(PageBreak())
        
        # ===== ACKNOWLEDGEMENTS =====
        ack = template.generate_acknowledgement_page()
        story.append(Paragraph("ACKNOWLEDGEMENTS", heading_style))
        story.append(Spacer(1, 20))
        for para in ack["content"].split("\n\n"):
            if para.strip():
                clean_para = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(clean_para, body_style))
                story.append(Spacer(1, 12))
        story.append(PageBreak())
        
        # ===== ABSTRACT =====
        story.append(Paragraph("ABSTRACT", heading_style))
        story.append(Spacer(1, 20))
        
        sections = proposal.get("full_sections", [])
        abstract_content = ""
        for section in sections:
            if "introduction" in section.get("title", "").lower():
                words = section.get("content", "").split()[:300]
                abstract_content = " ".join(words) + "..."
                break
        
        if abstract_content:
            clean = abstract_content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(clean, body_style))
        story.append(PageBreak())
        
        # ===== MAIN CONTENT =====
        for i, section in enumerate(sections, 1):
            title = section.get("title", f"Section {i}")
            content = section.get("content", "")
            
            story.append(Paragraph(title, heading_style))
            
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    clean_para = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(clean_para, body_style))
                    story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{proposal_id[:8]}_academic_proposal.pdf"'
            }
        )
        
    except ImportError:
        raise HTTPException(status_code=501, detail="reportlab not installed")


# ============================================================================
# Formatting Controller Routes
# ============================================================================

@extended_router.post("/proposals/{proposal_id}/remove-markdown")
async def remove_markdown_formatting(proposal_id: str):
    """
    Remove all markdown artifacts from a proposal.
    
    This ensures the output appears fully human-authored.
    """
    try:
        from src.services.formatting_controller import get_formatting_controller
        
        # Get the proposal
        import sys
        main_module = sys.modules.get('src.api.main')
        if main_module:
            completed_proposals = getattr(main_module, 'completed_proposals', {})
            proposal = completed_proposals.get(proposal_id)
        else:
            proposal = None
        
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        formatter = get_formatting_controller()
        formatted = formatter.format_proposal(proposal)
        
        # Update the stored proposal
        if main_module:
            completed_proposals[proposal_id] = formatted
        
        return {
            "success": True,
            "proposal_id": proposal_id,
            "original_word_count": proposal.get("word_count", 0),
            "formatted_word_count": formatted.get("word_count", 0),
            "message": "Markdown artifacts removed successfully",
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Formatting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@extended_router.post("/proposals/{proposal_id}/consolidate-references")
async def consolidate_references(proposal_id: str):
    """
    Consolidate all references into a single REFERENCES section.
    
    Removes inline reference lists from subsections.
    """
    try:
        from src.services.formatting_controller import get_formatting_controller, FormattingConfig
        
        # Get the proposal
        import sys
        main_module = sys.modules.get('src.api.main')
        if main_module:
            completed_proposals = getattr(main_module, 'completed_proposals', {})
            proposal = completed_proposals.get(proposal_id)
        else:
            proposal = None
        
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        config = FormattingConfig(
            consolidate_references=True,
            remove_inline_references=True,
            citation_style="harvard",
            use_et_al=True,
        )
        
        formatter = get_formatting_controller(config)
        formatted = formatter.format_proposal(proposal)
        
        # Update the stored proposal
        if main_module:
            completed_proposals[proposal_id] = formatted
        
        return {
            "success": True,
            "proposal_id": proposal_id,
            "message": "References consolidated successfully",
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Reference consolidation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
